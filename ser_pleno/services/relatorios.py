import json
import pandas as pd
import datetime
from docx import Document
from tkinter import filedialog
from config.db_config import get_db_connection

# Bibliotecas do ReportLab para PDF estruturado
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm

class ServicoRelatorio:
    
    def _execute_query(self, query, params=None, fetch=True, dictionary=True):
        """Método auxiliar para gestão de conexão"""
        conn = get_db_connection()
        try:
            cursor = conn.cursor(dictionary=dictionary)
            cursor.execute(query, params or ())
            if fetch:
                return cursor.fetchall()
            conn.commit()
            return cursor.lastrowid
        finally:
            conn.close()

    def listar_relatorios(self, tipo=None, data_inicio=None, pagina=1):
        query = "SELECT * FROM desktop_report WHERE 1=1"
        params = []
        
        if tipo and tipo != "Todos":
            query += " AND report_type = %s"
            params.append(tipo)
            
        offset = (pagina - 1) * 10
        query += " ORDER BY generated_at DESC LIMIT 10 OFFSET %s"
        params.append(offset)
        
        rows = self._execute_query(query, params)
        
        relatorios = [{
            'id': r['id'],
            'name': r['name'],
            'type': r['report_type'],
            'format': r['format'],
            'generated_at': r['generated_at'].strftime('%d/%m/%Y %H:%M') if r['generated_at'] else "--",
            'file_path': r['file_path']
        } for r in rows]
            
        return {"success": True, "data": {"reports": relatorios}}

    def obter_estatisticas(self):
        queries = {
            'students_total': "SELECT COUNT(*) as total FROM aluno",
            'appointments_total': "SELECT COUNT(*) as total FROM agendamento WHERE status = 'completed'",
            'interventions_total': "SELECT COUNT(*) as total FROM agendamento WHERE status = 'scheduled'",
            'screenings_total': "SELECT COUNT(*) as total FROM desktop_screening",
            'attendance_rate': "SELECT (COUNT(CASE WHEN status='completed' THEN 1 END) * 100 / NULLIF(COUNT(*), 0)) as total FROM agendamento"
        }
        
        resumo = {}
        for key, sql in queries.items():
            result = self._execute_query(sql)
            resumo[key] = int(result[0]['total'] or 0)
        
        return {"success": True, "data": {"summary": resumo}}

    def obter_dados_grafico(self):
        query = """
            SELECT DATE_FORMAT(data_hora, '%d/%m') as dia, COUNT(*) as total 
            FROM agendamento 
            WHERE data_hora >= DATE_SUB(NOW(), INTERVAL 7 DAY)
            GROUP BY dia 
            ORDER BY data_hora ASC
        """
        rows = self._execute_query(query)
        return {"success": True, "data": {"labels": [r['dia'] for r in rows], "values": [r['total'] for r in rows]}}

    def obter_tipos_relatorio(self):
        return [
            ('general', 'Relatório Geral'),
            ('student', 'Estudante'),
            ('appointments', 'Agendamentos'),
            ('interventions', 'Intervenções')
        ]

    # --- Exportação de Arquivos ---
    
    # --- Métodos de Busca para Exportação ---
    
    def exportar_estudantes(self, formato):
        query = "SELECT nome AS Nome, sala AS Sala, curso AS Curso FROM aluno ORDER BY nome ASC"
        dados = self._execute_query(query)
        return self._gerar_arquivo_fisico(dados, "Relatorio_Estudantes", formato)

    def exportar_agendamentos(self, formato):
        # Ajuste os nomes das colunas conforme seu banco real
        query = """
            SELECT a.data_hora AS Horario, al.nome AS Aluno, a.status AS Status 
            FROM agendamento a
            JOIN aluno al ON a.aluno_id = al.id 
            ORDER BY a.data_hora DESC
        """
        dados = self._execute_query(query)
        return self._gerar_arquivo_fisico(dados, "Relatorio_Agenda", formato)

    def exportar_triagens(self, formato):
        query = "SELECT aluno_nome AS Aluno, data_triagem AS Data, resultado AS Resultado FROM desktop_screening"
        dados = self._execute_query(query)
        return self._gerar_arquivo_fisico(dados, "Relatorio_Triagens", formato)

    def _gerar_arquivo_fisico(self, dados, nome_sugerido, formato):
        if not dados: return False

        extensoes = {
            "excel": (".xlsx", [("Excel", "*.xlsx")]),
            "pdf": (".pdf", [("PDF", "*.pdf")]),
            "word": (".docx", [("Word", "*.docx")])
        }
        
        ext, ft = extensoes.get(formato, (".pdf", [("PDF", "*.pdf")]))
        caminho = filedialog.asksaveasfilename(initialfile=nome_sugerido + ext, defaultextension=ext, filetypes=ft)
        if not caminho: return False

        try:
            if formato == "excel":
                pd.DataFrame(dados).to_excel(caminho, index=False)
            elif formato == "word":
                self._gerar_word_profissional(dados, caminho, nome_sugerido)
            elif formato == "pdf":
                self._gerar_pdf_profissional(dados, caminho, nome_sugerido)
            return True
        except Exception as e:
            print(f"Erro na exportação: {e}")
            return False

    def _gerar_word_profissional(self, dados, caminho, titulo):
        """Cria um Word com tabela real e editável"""
        doc = Document()
        doc.add_heading(titulo.replace("_", " "), 0)
        doc.add_paragraph(f"Gerado em: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        colunas = list(dados[0].keys())
        table = doc.add_table(rows=1, cols=len(colunas))
        table.style = 'Table Grid'
        
        # Cabeçalho
        for i, col_name in enumerate(colunas):
            table.rows[0].cells[i].text = col_name.upper()
        
        # Dados
        for item in dados:
            row_cells = table.add_row().cells
            for i, col_name in enumerate(colunas):
                row_cells[i].text = str(item.get(col_name, ""))
        
        doc.save(caminho)

    def _gerar_pdf_profissional(self, dados, caminho, titulo):
        """Cria um PDF organizado usando tabelas (Platypus)"""
        doc = SimpleDocTemplate(caminho, pagesize=A4, margin=2*cm)
        elements = []
        styles = getSampleStyleSheet()

        # Título e Data
        elements.append(Paragraph(f"<b>{titulo.replace('_', ' ').upper()}</b>", styles["Title"]))
        elements.append(Paragraph(f"Emitido em: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
        elements.append(Spacer(1, 12))

        # Preparar dados para a tabela
        headers = [list(dados[0].keys())]
        rows = [list(item.values()) for item in dados]
        data_table = headers + rows

        # Criar Tabela PDF
        t = Table(data_table, hAlign='LEFT', colWidths=[doc.width/len(headers[0])] * len(headers[0]))
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.beige])
        ]))
        
        elements.append(t)
        doc.build(elements)

    def excluir_relatorio(self, relatorio_id):
        """Executa a exclusão no banco de dados"""
        query = "DELETE FROM desktop_report WHERE id = %s"
        try:
            self._execute_query(query, (relatorio_id,), fetch=False)
            return {"success": True}
        except Exception as e:
            print(f"Erro ao excluir: {e}")
            return {"success": False, "error": str(e)}